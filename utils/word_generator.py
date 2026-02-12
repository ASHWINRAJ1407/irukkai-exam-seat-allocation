"""Word (DOCX) generation for seat allocation documents."""

from io import BytesIO

from docx import Document


def _add_header(document, exam_name: str, exam_date: str):
    """Common header with placeholders that admins can easily edit."""
    title = document.add_heading("EXAM SEATING ALLOCATION", level=1)
    title.alignment = 1  # center

    p = document.add_paragraph()
    run = p.add_run(f"Exam: {exam_name}    Date: {exam_date}")
    run.bold = True

    document.add_paragraph("Institution: ______________________________")
    document.add_paragraph("Invigilator: ______________________________")
    document.add_paragraph()  # blank line


def _merge_cells(table, col, start_row, end_row):
    """Merge cells in column col from start_row to end_row (inclusive).
    Clear text in merged cells so the value appears only once in the merged range.
    """
    if start_row >= end_row:
        return
    # Clear text in rows that will be merged (except the first) so the merged cell shows the value once
    for r in range(start_row + 1, end_row + 1):
        table.rows[r].cells[col].text = ''
    start_cell = table.rows[start_row].cells[col]
    end_cell = table.rows[end_row].cells[col]
    start_cell.merge(end_cell)


def create_overall_allocation_docx(allocations_by_hall, exam_name: str, exam_date: str) -> BytesIO:
    """
    Overall seating allocation in a single Word document.

    The generated file is intentionally simple so that admins can
    freely edit text, add institution name, invigilator name, etc.
    """
    document = Document()
    _add_header(document, exam_name, exam_date)

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Hall Number"
    hdr_cells[1].text = "Building Name"
    hdr_cells[2].text = "Floor Name"
    hdr_cells[3].text = "Department"
    hdr_cells[4].text = "No. Students"
    hdr_cells[5].text = "Roll Number Range"

    row_idx = 0

    for hall_info in allocations_by_hall:
        hall_num = hall_info.get("hall_number", "")
        building = hall_info.get("building_name", "") or ""
        floor = hall_info.get("floor", "") or ""

        for alloc in hall_info.get("allocations", []):
            dept = alloc.get("department_code") or alloc.get("dept_name") or ""
            count = alloc.get("count") or len(alloc.get("roll_numbers", [])) or 0

            rolls = alloc.get("roll_range", "")
            roll_numbers = alloc.get("roll_numbers", [])
            if not rolls and roll_numbers:
                rolls = (
                    f"{min(roll_numbers)} - {max(roll_numbers)}"
                    if len(roll_numbers) > 1
                    else str(roll_numbers[0])
                )

            row_idx += 1
            row_cells = table.add_row().cells
            row_cells[0].text = str(hall_num)
            row_cells[1].text = str(building)
            row_cells[2].text = str(floor)
            row_cells[3].text = str(dept)
            row_cells[4].text = str(count)
            row_cells[5].text = str(rolls)

    # Merge Hall Number, Building Name, Floor Name for consecutive rows of same hall
    if row_idx >= 1:
        current_hall = None
        start_row = 1
        for r in range(1, row_idx + 1):
            hall_val = table.rows[r].cells[0].text
            if hall_val != current_hall:
                if current_hall is not None and r - 1 > start_row:
                    for col in (0, 1, 2):
                        _merge_cells(table, col, start_row, r - 1)
                current_hall = hall_val
                start_row = r
        if row_idx > start_row:
            for col in (0, 1, 2):
                _merge_cells(table, col, start_row, row_idx)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

